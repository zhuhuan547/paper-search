# -*- coding: utf-8 -*-
"""生成简化版 prompt.docx — 个人使用的浏览器爬虫论文搜索工具方案"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def create_document():
    doc = Document()

    # Page Setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ========== 标题 ==========
    title = doc.add_heading('浏览器爬虫论文搜索工具 — 方案设计', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ========== 1. 工具定位 ==========
    doc.add_heading('1. 工具定位', level=1)
    doc.add_paragraph(
        '一个面向个人研究者使用的轻量级论文搜索工具。利用 Playwright Chrome Extension '
        '连接用户本地浏览器，复用浏览器中已登录的学术网站会话（Google Scholar、arXiv、'
        'DBLP、GitHub 等），避免单独处理登录认证，实现「即开即用」的论文爬取与筛选。'
    )

    doc.add_heading('1.1 核心思路', level=2)
    for item in [
        '用户正常使用自己的 Chrome 浏览器，已登录各大学术网站和 GitHub。',
        '工具通过 Playwright Chrome Extension 连接到当前浏览器，直接复用现有的登录态和 Cookie。',
        '在浏览器中自动操作页面（搜索、翻页、提取信息、访问GitHub验证代码等），和手动操作一模一样。',
        '所有操作在用户自己的浏览器中进行，不涉及额外账号或代理配置。',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    # ========== 2. 筛选条件设计 ==========
    doc.add_heading('2. 筛选条件设计', level=1)
    doc.add_paragraph('用户输入以下筛选条件，工具按条件在多个学术网站中搜索并筛选论文：')

    doc.add_heading('2.1 年份范围', level=2)
    doc.add_paragraph(
        '指定发表年份区间（如 2020-2024），支持单一年份或开区间（>=2021）。'
        '从论文元数据中提取年份进行过滤。'
    )

    doc.add_heading('2.2 论文任务类型', level=2)
    doc.add_paragraph(
        '指定模型解决的任务，如 image classification、object detection、'
        'semantic segmentation、visual question answering 等。通过论文标题和摘要中的'
        '关键词进行匹配。'
    )

    doc.add_heading('2.3 代码开源状态（含真伪验证）', level=2)
    doc.add_paragraph('分两步进行：')
    for item in [
        '第一步 — 找到代码链接：从论文页面（arXiv的comments、PapersWithCode链接、论文正文中的GitHub URL）提取代码仓库地址。',
        '第二步 — 验证代码真伪：通过 Playwright 在用户浏览器中打开 GitHub 仓库页面，检查：',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    for sub in [
        '文件列表中是否包含真正的代码文件（.py / .js / .ipynb 等），而不仅仅是 README.md + LICENSE',
        '代码文件数量和大概规模（通过页面元素判断即可，无需 clone）',
        '如果只有文档文件没有代码 → 标记为「假开源」，直接排除',
    ]:
        p = doc.add_paragraph(sub, style='List Bullet 2')

    doc.add_heading('2.4 模型模块的包含/排除', level=2)
    doc.add_paragraph(
        '用户指定模型包含或不包含的技术组件。例如：'
    )
    for item in [
        '不包含预训练编码器（如 ResNet、ViT、BERT 等预训练权重）',
        '包含对比学习（Contrastive Learning）',
        '包含 Attention 但不包含 Cross-Attention',
        '包含知识蒸馏',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        '实现方式：利用 LLM（如 GPT-4o 或 Claude）读取论文的 Method 章节，提取其中的技术方法列表，'
        '然后与用户的条件进行匹配。将每篇论文的模块标签缓存下来，避免重复调用 LLM。'
    )

    doc.add_heading('2.5 数据集筛选', level=2)
    doc.add_paragraph('用户可按以下维度筛选：')
    for item in [
        '数据集名称：如 Kvasir-SEG、ImageNet、COCO',
        '数据模态：图像 / 文本 / 视频 / 图文多模态',
        '数据领域：医学影像（内镜/CT/MRI）、自然图像、遥感 等',
    ]:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph(
        '实现方式：LLM 从论文中提取使用的数据集列表，与用户条件进行匹配。'
    )

    doc.add_heading('2.6 期刊/会议等级', level=2)
    doc.add_paragraph('支持按以下方式筛选：')
    for item in [
        '具体会议/期刊名：AAAI, CVPR, NeurIPS, ICCV, TPAMI 等',
        'CCF 等级：CCF-A / CCF-B / CCF-C',
        '也可仅看预印本（arXiv），不限期刊',
    ]:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph(
        '实现方式：维护一份本地的会议/期刊等级对照表（CCF推荐目录），直接查表匹配。'
    )

    doc.add_heading('2.7 其他实用筛选（可选）', level=2)
    for item in [
        '引用次数下限：从 Semantic Scholar / Google Scholar 页面获取被引次数',
        'GitHub Star 数下限：从 GitHub 页面获取',
        '论文类型：研究论文 / 综述 / 技术报告',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    # ========== 3. 数据源与爬取路径 ==========
    doc.add_heading('3. 数据源与爬取路径', level=1)

    doc.add_paragraph('利用 Playwright Chrome Extension 连接用户本地浏览器，按以下路径获取论文信息：')

    sources = [
        ('Semantic Scholar (主力)',
         [
             '直接在搜索框中输入关键词 → 获取论文列表（标题、作者、年份、引用数、摘要）',
             '有官方 API 可配合使用，速度快',
             '每篇论文页面上通常有 PDF 链接和代码链接',
         ]),
        ('arXiv',
         [
             '通过搜索获取预印本论文全文 PDF',
             'Comments 字段中常包含 GitHub 链接',
             '可获取完整 Method 章节供 LLM 分析',
         ]),
        ('DBLP',
         [
             '按会议/期刊名浏览论文列表',
             '用于确认论文的正式出版信息（会议名、年份、页码）',
             '配合 CCF 等级表判断论文等级',
         ]),
        ('GitHub',
         [
             '打开论文对应的代码仓库页面',
             '检查文件列表：是否有 .py / .js / .ipynb 等代码文件',
             '获取 Star 数、最后更新时间',
             '判断是否为「假开源」（只有 README.md 的空壳仓库）',
         ]),
        ('Google Scholar (辅助)',
         [
             '当其他源缺少引用次数时作为补充',
             '用户浏览器已登录，不会被反爬限制',
         ]),
    ]
    for name, details in sources:
        doc.add_heading(f'3.{sources.index((name, details))+1} {name}', level=2)
        for d in details:
            doc.add_paragraph(d, style='List Bullet')

    # ========== 4. 方案流程 ==========
    doc.add_heading('4. 方案流程', level=1)

    doc.add_paragraph('整个工具按以下步骤运行：')

    steps = [
        ('Step 1 — 用户输入条件',
         '用户通过一个简单的配置文件（JSON 或 YAML）或命令行参数，设定所有筛选条件，包括：'
         '年份范围、任务关键词、期刊等级要求、模块包含/排除规则、数据集要求、代码开源要求等。'),
        ('Step 2 — 连接浏览器',
         '工具通过 Playwright Chrome Extension 连接到用户当前打开的 Chrome 浏览器。'
         '无需任何登录操作，直接复用浏览器中已有的各大学术网站和 GitHub 的登录会话。'),
        ('Step 3 — 多源搜索',
         '按用户设定的条件，依次在各数据源中搜索：'
         '\n  - 在 Semantic Scholar / DBLP 中按关键词+年份搜索论文列表'
         '\n  - 在 arXiv 中获取预印本全文'
         '\n  - 提取每篇论文的元数据（标题、作者、摘要、年份、会议/期刊）'),
        ('Step 4 — 初筛',
         '根据年份、期刊等级、任务匹配等基础条件进行第一轮筛选，过滤明显不符合的论文。'),
        ('Step 5 — 代码验证',
         '对初筛通过的论文，逐一在浏览器中打开其 GitHub 仓库页面：'
         '\n  - 检查文件列表中是否包含代码文件'
         '\n  - 排除只有 README.md 的空壳仓库'
         '\n  - 记录 Star 数、代码规模等信息'),
        ('Step 6 — LLM 深度分析',
         '对经过代码验证的论文，利用 LLM（GPT-4o / Claude）进行深度分析：'
         '\n  - 读取论文的 Method 章节（从 arXiv PDF 或网页摘要获取）'
         '\n  - 提取模型的模块组成（技术标签列表）'
         '\n  - 提取使用的数据集列表'
         '\n  - 与用户的模块条件、数据集条件进行匹配'),
        ('Step 7 — 结果输出',
         '将最终匹配的论文以结构化方式输出：'
         '\n  - 列表展示：标题、作者、年份、会议/期刊、引用数、代码状态、匹配的模块标签'
         '\n  - 导出格式：Markdown 表格 / CSV / BibTeX'
         '\n  - 可附带简要说明：为什么该论文符合/不符合条件'),
    ]

    for step_title, step_desc in steps:
        doc.add_heading(step_title, level=2)
        doc.add_paragraph(step_desc)

    # ========== 5. 代码真伪验证流程 ==========
    doc.add_heading('5. 代码验证流程（重点）', level=1)
    doc.add_paragraph(
        '这是本工具的核心功能之一。很多论文声称开源代码，但实际上 GitHub 仓库中只有一个 '
        'README.md 文件，没有真正的代码。验证流程如下：'
    )

    verify_steps = [
        '① 从论文信息中提取 GitHub URL（arXiv comments、论文正文链接、PapersWithCode 等来源）。',
        '② 在用户浏览器中打开该 GitHub 仓库页面。',
        '③ 查看仓库的文件列表（GitHub 网页上的文件树），统计文件类型：',
    ]
    for v in verify_steps:
        doc.add_paragraph(v, style='List Bullet')

    doc.add_paragraph('判定逻辑：', style='List Bullet')
    criteria = [
        '若文件列表中包含 .py / .js / .ts / .ipynb / .R / .cpp 等编程语言文件 → 通过，标记为「真开源」',
        '若文件列表中只有 README.md / LICENSE / .gitignore 等非代码文件 → 不通过，标记为「假开源」，排除',
        '若仓库文件总数 < 5 且无编程语言文件 → 不通过，排除',
        '若仓库是纯 Fork（与原仓库完全一致，无新增提交）→ 标记但暂不排除（用户可选择是否排除）',
    ]
    for c in criteria:
        doc.add_paragraph(c, style='List Bullet 2')

    doc.add_paragraph(
        '整个过程通过 Playwright 在浏览器中操作 GitHub 网页即可完成，不需要 git clone，快速且不占用本地磁盘。'
    )

    # ========== 6. 简化技术方案 ==========
    doc.add_heading('6. 简化技术方案', level=1)
    doc.add_paragraph('不搞复杂的微服务架构，就是一个本地运行的 Python 脚本：')

    tech = [
        ('浏览器自动化', 'Playwright + Playwright Chrome Extension\n（连接用户本地 Chrome，复用登录态）'),
        ('后端逻辑', 'Python 脚本，本地运行'),
        ('LLM 调用', 'OpenAI API (GPT-4o) 或 Claude API，用于论文内容分析和模块提取'),
        ('数据存储', '本地 JSON 文件 或 SQLite（轻量，无需安装数据库服务）'),
        ('结果输出', 'Markdown 表格 / CSV 文件 / 终端打印'),
        ('配置方式', 'JSON 配置文件或命令行参数'),
    ]
    for name, desc in tech:
        doc.add_heading(name, level=3)
        doc.add_paragraph(desc)

    # ========== 7. 用户使用示例 ==========
    doc.add_heading('7. 使用示例', level=1)

    doc.add_heading('7.1 配置文件示例 (search_config.json)', level=2)
    config_example = '''{
  "year_range": { "min": 2020, "max": 2024 },
  "task_keywords": ["medical image segmentation", "polyp detection"],
  "venue": {
    "ccf_rank": ["A"],
    "specific": ["MICCAI", "CVPR", "AAAI"]
  },
  "code_required": true,
  "module_rules": {
    "exclude": ["pretrained encoder", "ImageNet pretrained"],
    "include": ["attention mechanism"]
  },
  "dataset_rules": {
    "include_domains": ["endoscopy", "medical imaging"],
    "include_names": ["Kvasir-SEG"]
  },
  "output": {
    "format": "markdown",
    "max_results": 50
  }
}'''

    p = doc.add_paragraph()
    run = p.add_run(config_example)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)

    doc.add_heading('7.2 命令行使用', level=2)
    p = doc.add_paragraph()
    run = p.add_run('python search_papers.py --config search_config.json')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph(
        '运行后，工具自动连接用户的 Chrome 浏览器，在各学术网站中搜索，筛选符合条件的论文，'
        '验证代码仓库，最终输出一份 Markdown 格式的结果列表。'
    )

    # ========== 8. 注意事项 ==========
    doc.add_heading('8. 注意事项', level=1)
    for item in [
        '速度控制: 在页面操作之间加入合理延迟（2-5 秒），避免对网站造成压力或触发反爬。',
        'LLM 成本: 每篇论文的 Method 分析约消耗 1K-3K token，建议先用标题/摘要初筛，减少需要 LLM 分析的论文数量。',
        '结果缓存: 已分析过的论文结果缓存到本地 JSON/SQLite，下次查询时直接复用，节省时间和 LLM 费用。',
        '合规使用: 本工具仅供个人学术研究使用，请遵守各网站的 robots.txt 和服务条款。',
        '浏览器保持打开: 使用期间 Chrome 浏览器需要保持运行，工具通过 Chrome Extension 与浏览器通信。',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— 文档结束 —')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)

    output_path = r'D:\Claude_project\search_paper\prompt_v2.docx'
    doc.save(output_path)
    print(f'Done: {output_path}')


if __name__ == '__main__':
    create_document()
